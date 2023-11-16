import streamlit as st
from docx import Document
import base64
from openai import OpenAI
from prompts import PROBLEM_STATEMENT_PROMPT, NEED_FOR_INNOVATION_PROMPT, INNOVATION_NAME_PROMPT, INNOVATION_DESCRIPTION_PROMPT, GLOBAL_TEMPERATURE

# Set the page config
st.set_page_config(
    page_title="Generador de Ideas de Innovación",
    page_icon="🚀"
)

# Custom CSS to hide default footer and show custom footer
custom_css = """
<style>
    footer {
        visibility: hidden;
    }
    footer:after {
        content:'Desarrollado por MEIK LABS'; 
        visibility: visible;
        display: block;
        position: relative;
        padding: 5px;
        top: 2px;
    }
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# Load secrets
api_key = st.secrets["general"]["OPENAI_API_KEY"]
temperature = st.secrets.get("temperature", 0.9)
model_name = st.secrets.get("fast_llm_model", "gpt-3.5-turbo")

client = OpenAI(api_key=api_key)

# Initialize OpenAI API
#openai.api_key = api_key

st.title("🚀 Generador de Ideas de Innovación")

def make_downloadable_link(file_path, file_name, download_name):
    with open(file_path, "rb") as f:
        bytes_data = f.read()
        b64 = base64.b64encode(bytes_data).decode()
        return f'<a href="data:application/octet-stream;base64,{b64}" download="{download_name}">{file_name}</a>'

def refine_input(input_text, max_tokens=150, temperature=0.8):
    try:
        # OpenAI Chat API call
        response = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "You are an Innovation Expert, helping to present an innovation idea."},
                {"role": "user", "content": input_text}
            ],
            max_tokens=max_tokens,
            temperature=temperature
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.warning(f"OpenAI API Error: {str(e)}")
        return input_text
    except Exception as e:
        st.warning(f"OpenAI API Error: {str(e)}")
        return input_text



with st.form(key='innovation_form'):
    # Input fields
    innovation_name = st.text_input("💡 Nombre de la Idea", max_chars=50)
    problem_statement = st.text_area("🚧 Descripción del Problema", max_chars=500, help="Describe your problem statement. Address only one problem at once.")
    need_for_innovation = st.text_area("👩‍🚀¿A quiénes les afecta el problema?", max_chars=200)
    innovation_description = st.text_area("🏆 Describe la solución", max_chars=300)
    references = st.text_area("🔍 Referencias (links)", help="Provide URLs separated by newline")
            
    # Submit button
    submit = st.form_submit_button("Enviar")

# When form is submitted
if submit:
    # Process the references
    processed_references = references.split('\n')

    with st.spinner('Generando el reporte...'):
        # Create a new Document instance
        doc = Document()
        doc.add_heading('Idea de Innovación', level=1)

        # Use ChatGPT refinement function for the user inputs with prompts
        doc.add_heading('Nombre de la Idea:', level=1)
        doc.add_paragraph(refine_input(INNOVATION_NAME_PROMPT['prompt'] + innovation_name, INNOVATION_NAME_PROMPT['max_tokens'], GLOBAL_TEMPERATURE))
        
        doc.add_heading('Descripción del Problema:', level=2)
        doc.add_paragraph(refine_input(PROBLEM_STATEMENT_PROMPT['prompt'] + problem_statement, PROBLEM_STATEMENT_PROMPT['max_tokens'], GLOBAL_TEMPERATURE))

        doc.add_heading('¿A quiénes les afecta el problema?:', level=2)
        doc.add_paragraph(refine_input(NEED_FOR_INNOVATION_PROMPT['prompt'] + need_for_innovation, NEED_FOR_INNOVATION_PROMPT['max_tokens'], GLOBAL_TEMPERATURE))
        
        doc.add_heading('Descripción de la solución:', level=2)
        doc.add_paragraph(refine_input(INNOVATION_DESCRIPTION_PROMPT['prompt'] + innovation_description, INNOVATION_DESCRIPTION_PROMPT['max_tokens'], GLOBAL_TEMPERATURE))

        doc.add_heading('Referencias:', level=2)
        for ref in processed_references:
            doc.add_paragraph(ref)

        
        # Save the report with the sanitized innovation_name as filename
        sanitized_filename = "".join(e for e in innovation_name if e.isalnum() or e == " ")  # Removing special characters for safety
        sanitized_filename = sanitized_filename.replace(" ", "_")  # Replacing spaces with underscores
        doc.save(f"{sanitized_filename}.docx")

    # After the spinner ends, display success message and provide the download link
    st.success('Reporte generado exitosamente!')
    download_link = make_downloadable_link(f"{sanitized_filename}.docx", "Descargar el reporte", f"{sanitized_filename}.docx")
    st.markdown(download_link, unsafe_allow_html=True)